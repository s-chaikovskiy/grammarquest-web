# -*- coding: utf-8 -*-
"""
Сборка документов олимпиады для версии 2.0.

Два документа собираются по оглавлению прошлогодней защиты: те же разделы,
та же нумерация, тот же порядок. Так учителю не нужно сверять структуру —
достаточно прочитать содержание.

Текст лежит в `техническое.py` и `руководство.py`, вёрстка — здесь.
Разделение то же, что и в презентации: правка текста не должна требовать
правки оформления.

    ./venv/bin/python сборка.py

Шрифт Times New Roman 12 — как принято в школьных работах. Заголовки
жирные того же кегля, что и в оригинале: 20 / 16 / 14.
"""
import json
import pathlib
import re
import subprocess
import sys
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = pathlib.Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

# Читалка PDF: постранично отдаёт текст. Своя, потому что pdftotext на этой
# машине нет, а ставить его ради одной операции незачем — PDFKit уже в системе.
ЧТЕЦ = HERE / "чтец-pdf"

BODY = "Times New Roman"
MONO = "Consolas"
INK = RGBColor(0x1B, 0x22, 0x2B)
GREY = RGBColor(0x55, 0x5D, 0x68)
# Тот же синий, что и в приложении (--accent-ink, oklch(0.43 0.15 250)).
# Цвет здесь ровно один и только у линеек под разделами: документ школьный,
# читать его будут на бумаге, и цветным текстом он выглядел бы рекламой.
АКЦЕНТ = RGBColor(0x00, 0x50, 0x9D)


def новый_документ(заголовок: str, подзаголовок: str, подпись: str = "") -> Document:
    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21), Cm(29.7)
    s.left_margin = s.right_margin = Cm(3)
    s.top_margin = s.bottom_margin = Cm(2)

    n = d.styles["Normal"]
    n.font.name = BODY
    n.font.size = Pt(12)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.15

    for имя, кегль in (("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 14)):
        st = d.styles[имя]
        st.font.name = BODY
        st.font.size = Pt(кегль)
        st.font.bold = True
        st.font.color.rgb = INK
        st.paragraph_format.space_before = Pt(16 if кегль > 14 else 12)
        st.paragraph_format.space_after = Pt(6)

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(заголовок)
    r.font.size, r.font.bold, r.font.name = Pt(22), True, BODY

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(подзаголовок)
    r.font.size, r.font.name = Pt(13), BODY
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(18)

    if подпись:
        _колонтитул(d, подпись)
    return d


def _линейка(параграф, цвет=АКЦЕНТ, толщина=8):
    """Тонкая черта под абзацем: в python-docx границ нет, только XML."""
    pPr = параграф._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(толщина))       # восьмые доли пункта
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), f"{цвет}")
    borders.append(bottom)
    pPr.append(borders)


def _поле(параграф, инструкция: str, запасное: str = "1"):
    """Поле Word (например PAGE): номер страницы считает сам редактор."""
    r = параграф.add_run()
    # Оформление задаём до вставки поля: иначе w:rPr встанет после узлов
    # поля и номер напечатается стилем по умолчанию — крупнее подписи.
    r.font.name, r.font.size = BODY, Pt(9.5)
    r.font.color.rgb = GREY
    начало = OxmlElement("w:fldChar")
    начало.set(qn("w:fldCharType"), "begin")
    текст = OxmlElement("w:instrText")
    текст.set(qn("xml:space"), "preserve")
    текст.text = инструкция
    разделитель = OxmlElement("w:fldChar")
    разделитель.set(qn("w:fldCharType"), "separate")
    значение = OxmlElement("w:t")
    значение.text = запасное
    конец = OxmlElement("w:fldChar")
    конец.set(qn("w:fldCharType"), "end")
    for узел in (начало, текст, разделитель, значение, конец):
        r._r.append(узел)
    return r


def _колонтитул(d: Document, подпись: str):
    """Нижний колонтитул: слева название работы, справа номер страницы."""
    # У стиля «Footer» своя табуляция по центру страницы. Табуляции абзаца
    # не заменяют стилевые, а складываются с ними, поэтому первый \t уводил
    # номер в середину строки. Чистим у стиля, а не у абзаца.
    подвал = d.styles["Footer"]
    подвал.paragraph_format.tab_stops.clear_all()
    # Оформление задаём стилю, а не рунам: номер страницы — поле, и редактор
    # пересчитывает его своим руном. Оформление рунa при этом теряется,
    # стилевое — нет.
    подвал.font.name, подвал.font.size = BODY, Pt(9.5)
    подвал.font.color.rgb = GREY

    ф = d.sections[0].footer
    p = ф.paragraphs[0] if ф.paragraphs else ф.add_paragraph()
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(f"{подпись}\t")
    r.font.name, r.font.size = BODY, Pt(9.5)
    r.font.color.rgb = GREY
    _поле(p, " PAGE ")


def _оглавление(d: Document, пункты, страницы=None):
    """
    «Мазмұны»: только разделы верхнего уровня — их дюжина, и лист остаётся
    обозримым. Подразделов полсотни, они превратили бы содержание в простыню.

    Номера страниц ставятся вторым проходом: первый раз документ собирается
    с прочерками, обращается в PDF, и оттуда становится видно, где какой
    раздел лёг. Число строк при этом не меняется, поэтому вёрстка второго
    прохода совпадает с первым — номера получаются верные.
    """
    з = d.add_paragraph()
    r = з.add_run("Мазмұны")
    r.font.size, r.font.bold, r.font.name = Pt(16), True, BODY
    r.font.color.rgb = INK
    з.paragraph_format.space_after = Pt(10)
    _линейка(з)

    for i, пункт in enumerate(пункты):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        номер = "—" if страницы is None else str(страницы[i])
        r = p.add_run(f"{пункт}\t{номер}")
        r.font.name, r.font.size = BODY, Pt(12)

    d.add_page_break()


def _маркер(d, текст, уровень=0):
    p = d.add_paragraph(текст, style="List Bullet" if уровень == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.name, r.font.size = BODY, Pt(12)
    return p


def _код(d, текст):
    for строка in текст.strip("\n").split("\n"):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(строка if строка.strip() else " ")
        r.font.name, r.font.size = MONO, Pt(9.5)
        r.font.color.rgb = INK


def _таблица(d, строки):
    t = d.add_table(rows=len(строки), cols=len(строки[0]))
    t.style = "Table Grid"
    for i, ряд in enumerate(строки):
        for j, ячейка in enumerate(ряд):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(ячейка))
            r.font.name, r.font.size = BODY, Pt(10.5)
            if i == 0:
                r.font.bold = True
    d.add_paragraph().paragraph_format.space_after = Pt(4)


def _картинка(d, путь, подпись):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(путь), width=Cm(7.2))
    c = d.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(подпись)
    r.font.size, r.font.italic, r.font.name = Pt(10), True, BODY
    r.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(12)


def собрать(блоки, d: Document, база: pathlib.Path):
    """Блок — кортеж («вид», …). Виды: h1 h2 h3 p b b2 code table img."""
    for блок in блоки:
        вид, данные = блок[0], блок[1]
        if вид in ("h1", "h2", "h3"):
            з = d.add_heading(данные, level=int(вид[1]))
            if вид == "h1":
                _линейка(з)
        elif вид == "p":
            p = d.add_paragraph(данные)
            for r in p.runs:
                r.font.name, r.font.size = BODY, Pt(12)
        elif вид == "b":
            for пункт in данные:
                _маркер(d, пункт)
        elif вид == "b2":
            for пункт in данные:
                _маркер(d, пункт, 1)
        elif вид == "code":
            _код(d, данные)
        elif вид == "table":
            _таблица(d, данные)
        elif вид == "img":
            путь = база / данные
            if путь.exists():
                _картинка(d, путь, блок[2])
            else:
                print(f"  ! нет картинки: {путь.name}")
        else:
            raise ValueError(f"неизвестный вид блока: {вид}")
    return d


SOFFICE = "/usr/local/bin/soffice"


def собрать_чтец():
    """Компилируем читалку, если её ещё нет: в репозитории лежит только код."""
    if ЧТЕЦ.exists():
        return
    subprocess.run(["swiftc", "-O", str(ЧТЕЦ.with_suffix(".swift")), "-o", str(ЧТЕЦ)],
                   check=True, capture_output=True)


def в_pdf(docx: pathlib.Path) -> pathlib.Path:
    """DOCX → PDF через LibreOffice: он же считает поля и номера страниц."""
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", str(docx.parent), str(docx)],
        check=True, capture_output=True,
    )
    return docx.with_suffix(".pdf")


def страницы_разделов(pdf: pathlib.Path, разделы) -> list[int]:
    """
    Ищем каждый раздел в тексте PDF постранично. Первая страница пропускается:
    там оглавление, и все названия встречаются на ней ещё раз.
    """
    вывод = subprocess.run([str(ЧТЕЦ), str(pdf)], check=True, capture_output=True).stdout
    # PDFKit отдаёт текст в разложенном виде: «й» приходит как «и» плюс
    # значок сверху отдельным символом. Без сборки обратно в NFC половина
    # казахских заголовков не находится, хотя на экране выглядит так же.
    плоско = lambda т: re.sub(r"\s+", "", unicodedata.normalize("NFC", т))
    листы = [плоско(лист) for лист in json.loads(вывод)]
    номера = []
    for раздел in разделы:
        игла = плоско(раздел)
        найден = next((i + 1 for i, лист in enumerate(листы[1:], start=1) if игла in лист), None)
        if найден is None:
            raise SystemExit(f"не нашёл раздел в PDF: {раздел}")
        номера.append(найден)
    return номера


def построить(модуль, скрины: pathlib.Path, страницы=None) -> pathlib.Path:
    разделы = [b[1] for b in модуль.БЛОКИ if b[0] == "h1"]
    d = новый_документ(модуль.ЗАГОЛОВОК, модуль.ПОДЗАГОЛОВОК, модуль.КОЛОНТИТУЛ)
    _оглавление(d, разделы, страницы)
    собрать(модуль.БЛОКИ, d, скрины)
    путь = HERE / модуль.ФАЙЛ
    d.save(путь)
    return путь


def main():
    import руководство
    import техническое

    собрать_чтец()
    скрины = HERE / "скриншоты"
    for модуль in (техническое, руководство):
        разделы = [b[1] for b in модуль.БЛОКИ if b[0] == "h1"]

        # Первый проход — с прочерками вместо номеров, только чтобы узнать,
        # на какой лист лёг каждый раздел.
        pdf = в_pdf(построить(модуль, скрины))
        номера = страницы_разделов(pdf, разделы)

        # Второй проход — с настоящими номерами. Строк столько же, поэтому
        # вёрстка не сдвигается; третьим проходом это проверяем.
        pdf = в_pdf(построить(модуль, скрины, номера))
        проверка = страницы_разделов(pdf, разделы)
        if проверка != номера:
            raise SystemExit(
                f"{модуль.ФАЙЛ}: оглавление разъехалось, {номера} → {проверка}")

        знаков = sum(len(b[1]) if isinstance(b[1], str)
                     else sum(len(str(x)) for x in b[1]) for b in модуль.БЛОКИ)
        print(f"  собрано: {модуль.ФАЙЛ}  ({len(модуль.БЛОКИ)} блоков, ~{знаков} знаков, "
              f"{len(разделы)} разделов, {проверка[-1]}+ стр.)", flush=True)

    # Сверка чисел идёт сразу за сборкой, а не по памяти: в документах
    # 26 сабақ, 623 жазба, 70 автотест — и разъезжаются они тихо.
    # Так уже случилось с весом приложения: в тексте стояло 202 КБ
    # при фактических 210.
    print()
    subprocess.run([sys.executable, str(HERE / "сверка-чисел.py")], check=False)


if __name__ == "__main__":
    main()
